
import mecanico_linhas as ml
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import time
import xml.etree.ElementTree as ET
from matplotlib.patches import Polygon
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Cm
from docx.shared import Inches
from docx.enum.table import  WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Registre o tempo de início
tempo_inicial = time.time()
numero_documento = "Catu"
nome_LT = "LD 69 kV Catu - Entre Rios"

deflexao = 0
deflexao_bissetriz = 0
rugosidade = "B"
temperatura = {
    "coincidente": 16,
    "EDS": 21,
    "minima": 4,
    }

qtd_cabo_fase = 1
qtd_cabo_pr = 1
poste = "DT" #DT OU R , PARA TORRES DEIXAR SEM NADA
Carga_inicial = 1500



terreno = "B"
altitude = 100
cabo = "Linnet"
cabo_pr1 = "5/16'"
cabo_pr2 = "Penguin"
vento = {"projeto":23.62,"extremo":38.3 }
periodo_retorno = 150

tracao_eds = {"vante": 1400,"re": 1400}

tracao_eds_pr2 = {"vante": 400, "re": 400}
tracao_eds_pr = {"vante": 400, "re": 400}

vao = {"vante":  250 ,"re": 250}

vao_de_peso =  {"min":50,"max":300 }

vao_regulador = {"vante":  250 ,"re":250}


isolador = {"comprimento": 1.5,"area": 0.59, "peso": 50 }

nome_estrutura = " 0/2"
altura_estrutura = 21
geometria_estrutura = {
    "Fase_1": {"H": 3.8, "X": -1.9, "Tipo": "Suspensao", "cabo": cabo, "qtd": qtd_cabo_fase},
    "Fase_2": {"H": 2.04, "X": 1.8, "Tipo": "Suspensao", "cabo": cabo, "qtd": qtd_cabo_fase},
    "Fase_3": {"H": 4.73, "X": 1.8, "Tipo": "Suspensao", "cabo": cabo, "qtd": qtd_cabo_fase},
    "Cabo_PR_1": {"H": 0.2, "X": 0, "Tipo": "Suspensao", "cabo": cabo_pr1, "qtd": qtd_cabo_pr},
}

peso_homens = 200

dt_tipo = { 
    1500:3,
    2000:6,
    2500:6,
    3000:6,
    4000:9,  
    }

r_tipo = {
    1500:10,
    2000:10,
    2500:16,
    3000:18,
    4000:23
    }

dt_dim = {
    3:{"Face_A": 224+5, "Face_B": 170+5 },
    6:{"Face_A": 308+5, "Face_B": 230+5 },
    9:{"Face_A": 392+5, "Face_B": 290+5 },
    }

r_dim = {
    10:{"Face_A": 280+5, "Face_B": 280+5 },
    16:{"Face_A": 370+5, "Face_B": 370+5 },
    18:{"Face_A": 400+5, "Face_B": 400+5 },
    23:{"Face_A": 475+5, "Face_B": 475+5 },
    }

L_tronco  = altura_estrutura/4
incremento_A = 28*L_tronco
incremento_B = 20*L_tronco


tronco = {1:dt_dim[dt_tipo[Carga_inicial]]}



con_a = 28
con_b = 20

if poste =="R": 
    tronco = {1:r_dim[r_tipo[Carga_inicial]]}
    con_a = 15
    con_b = 15


for i in range(2, 2 + 4):
    tronco[i] = {'Face_A': tronco[i-1]['Face_A'] + incremento_A, 'Face_B': tronco[i-1]['Face_B'] + incremento_B}
    tronco[i-1]["area_A"] = (tronco[i-1]['Face_A'] +tronco[i]['Face_A'])* L_tronco/2000
    tronco[i-1]["area_B"] = (tronco[i-1]['Face_B'] +tronco[i]['Face_B'])* L_tronco/2000
    tronco[i-1]["centroide"] = round((L_tronco / 3) * (
        (tronco[i - 1]["Face_A"]*2 + tronco[i]["Face_A"]) /
        (tronco[i - 1]["Face_A"] + tronco[i]["Face_A"])) + (5-i)*L_tronco,2) 
    
    tronco[i - 1]["gt"] = round(ml.GT(tronco[i-1]["centroide"], rugosidade),2)
    tronco[i-1]["secaoA"] = round( (tronco[1]["Face_A"] + (altura_estrutura - tronco[i-1]["centroide"])*con_a)/1000,2)
    tronco[i-1]["secaoB"] = round((tronco[1]["Face_B"] + (altura_estrutura - tronco[i-1]["centroide"])*con_b)/1000,2)



dbCabos = dados_excel = pd.read_excel("_db/cabos.xlsx")
dadosCabo = dbCabos[dbCabos['Cabo'] == cabo]
dados_pr1 = dbCabos[dbCabos['Cabo'] == cabo_pr1]
dados_pr2 = dbCabos[dbCabos['Cabo'] == cabo_pr2]
tracao_min = {
    "vante": ml.mudanca_estado(
        float(dadosCabo["E"].iloc[0]),
        float(dadosCabo["S"].iloc[0]),
        float(vao["vante"]),
        float(dadosCabo["COEF"].iloc[0]),
        float(dadosCabo["Peso (kgf/m)"].iloc[0]),
        float(dadosCabo["Peso (kgf/m)"].iloc[0]),
        float(temperatura["EDS"]),
        float(temperatura["minima"]),
        float(tracao_eds["vante"]),
    ),
    
    "re":  ml.mudanca_estado(
        float(dadosCabo["E"].iloc[0]),
        float(dadosCabo["S"].iloc[0]),
        float(vao["re"]),
        float(dadosCabo["COEF"].iloc[0]),
        float(dadosCabo["Peso (kgf/m)"].iloc[0]),
        float(dadosCabo["Peso (kgf/m)"].iloc[0]),
        float(temperatura["EDS"]),
        float(temperatura["minima"]),
        float(tracao_eds["re"]),
    ),
}

Cxposte = 2 ## mudar para circular

if poste == "R":
    Cxposte = 1.2
#hipoteses
hipoteses = {
    1: {'Temperatura': temperatura["coincidente"], 'Vento': vento["projeto"], 'Ângulo': 90, "nome": "Vento máximo Transversal"},
    2: {'Temperatura': temperatura["coincidente"], 'Vento': vento["projeto"], 'Ângulo': 0, "nome": "Vento máximo Longitudinal"},
    3: {'Temperatura': temperatura["coincidente"], 'Vento': vento["projeto"], 'Ângulo': 45, "nome": "Vento máximo a 45 graus"},
    4: {'Temperatura': temperatura["coincidente"], 'Vento': vento["projeto"], 'Ângulo': 60, "nome": "Vento máximo a 60 graus"},
    5: {'Temperatura': temperatura["coincidente"], 'Vento': vento["projeto"], 'Ângulo': 75, "nome": "Vento máximo a 75 graus"},
    6: {'Temperatura': temperatura["coincidente"], 'Vento': vento["extremo"], 'Ângulo': 90, "nome": "Vento de alta intensidade Transversal"},
    #7: {'Temperatura': temperatura["coincidente"], 'Vento': vento["extremo"], 'Ângulo': 0, "nome": "Vento de alta intensidade Longitudinal"},
    8: {'Temperatura': temperatura["coincidente"], 'Vento': vento["extremo"], 'Ângulo': 45, "nome": "Vento de alta intensidade a 45 graus"},
    9: {'Temperatura': temperatura["coincidente"], 'Vento': vento["extremo"], 'Ângulo': 60, "nome": "Vento de alta intensidade a 60 graus"},
    10: {'Temperatura': temperatura["coincidente"], 'Vento': vento["extremo"], 'Ângulo': 75, "nome": "Vento de alta intensidade a 75 graus"},
    11: {'Temperatura': temperatura["EDS"], 'Vento': 0, 'Ângulo': 0, "nome": "Construção/Manutenção"},                           # construcao
    12: {'Temperatura': temperatura["EDS"], 'Vento': 0, 'Ângulo': 0, "nome": "Cabo condutor rompido a vante"},                           #cabo rompido vante
    #13: {'Temperatura': temperatura["EDS"], 'Vento': 0, 'Ângulo': 0,"nome": "Cabo condutor rompido a ré"},                           #cabo rompido re
    14: {'Temperatura': temperatura["EDS"], 'Vento': 0, 'Ângulo': 0,"nome": "PR rompido a vante"},                           #pr rompido vante
    15: {'Temperatura': temperatura["EDS"], 'Vento': 0, 'Ângulo': 0,"nome": "PR rompido a ré"},                           #pr rompido re 
    #16: {'Temperatura': temperatura["EDS"], 'Vento': 0, 'Ângulo': 0,"nome": "Contenção em cascata"},                           #Contenção em cascata
    17: {'Temperatura': temperatura["minima"], 'Vento': 0, 'Ângulo': 0,"nome": "Temperatura mínima"},
    #18: {'Temperatura': temperatura["coincidente"], 'Vento': vento["projeto"], 'Ângulo': 90, "nome": "Terminal"},
}

fator_sobrecarga = {
    "Ruptura de cabo":  {"T": 1, "L": 1, "V_normal": 1.15, "V_reduzido": 1.15 if vao_de_peso["min"] < 0 else 0.87},
    "Vento de alta intensidade":  {"T": 1, "L": 1, "V_normal": 1.15, "V_reduzido": 1.15 if vao_de_peso["min"] < 0 else 0.87},
    "Vento Máximo":  {"T": 1.15, "L": 1.15, "V_normal": 1.15, "V_reduzido": 1.15 if vao_de_peso["min"] < 0 else 0.87},
    "Contenção em cascata":  {"T": 1, "L": 1, "V_normal": 1.15, "V_reduzido": 1.15 if vao_de_peso["min"] < 0 else 0.87},
    "Construção/Manutenção": {"T": 1.5, "L":1.5, "V_normal": 1.5, "V_reduzido": 1.5 if vao_de_peso["min"] < 0 else 0.67 },
     }


massa = ml.massa_do_ar(temperatura["coincidente"], altitude)

pressoes_referencia = {
    key: ml.pressao(massa, value) for key, value in vento.items()
}
        
# CORRIGINDO A ALTURA DAS ESTRUTURAS
for chave, fase in geometria_estrutura.items():
     fase["H"] = altura_estrutura - fase["H"]
     peso = float(dbCabos.loc[dbCabos['Cabo'] == fase["cabo"], "Peso (kgf/m)"].iloc[0])
     fase["H_re"] =  round(fase["H"] - (peso*vao["re"]**2)/(12*tracao_min["re"]),2)
     fase["H_vante"] =  round(fase["H"] - (peso*vao["vante"]**2)/(12*tracao_min["vante"]),2)
     if fase["Tipo"] == "Suspensão":
         fase["H_re"] = fase["H_re"]-isolador["comprimento"]
         fase["H_vante"] = fase["H_vante"]-isolador["comprimento"]
     fase["gc_re"] = ml.GC(fase["H_re"], rugosidade)
     fase["gc_vante"] = ml.GC(fase["H_vante"], rugosidade)
     
gl = {"vante": ml.GL(vao["vante"]),
      "re": ml.GL(vao["re"]),
      }    


#CRIANDO A GEOMETRIA DO ISOLADOR
geometria_isolador = {}

for chave, fase in geometria_estrutura.items():
    if chave.startswith("Fase"):
        # Se a chave for uma fase, crie a geometria do isolador
        isolador_geometria = {
            "H": fase["H"] - isolador["comprimento"] / 2 if fase["Tipo"] == "Suspensão" else fase["H"],
            "X": fase["X"]
        }
        geometria_isolador[chave] = isolador_geometria
        
        
                    
for chave, fase in geometria_isolador.items():
    # Calcule o valor de gt para a altura H na fase e um tipo de terreno (você deve substituir 'Terreno' pelo valor real)
    gt = ml.GT(fase["H"], terreno)
    
    # Adicione o valor de gt à geometria do isolador
    fase["gt"] = gt

for fase, valores in geometria_isolador.items():
    valores['pre_pro'] = round(valores['gt'] *pressoes_referencia["projeto"]*1.2,2)
    valores['pre_ext'] =round( valores['gt'] *pressoes_referencia["extremo"]*1.2,2)



# Inicialize o dicionário para armazenar as forças para cada hipótese e fase
forca_isolador = {}
pressao_vento_cabo = {}
tracao_cabo = {}
tracao_cabo1 = {}
pressao_ven_cabo = {}
arvore_carga = {}
arvore_carga_poste = {}
forca_poste =  {}
Cx_cabo = 1
pressao_tronco = {}



# Loop sobre as hipóteses em hipoteses
for hipotese, parametros_hipotese in hipoteses.items():
    # Inicialize o dicionário para armazenar as forças para cada fase
    forcas_por_fase = {}
    temp = parametros_hipotese['Temperatura']
    vento_1= parametros_hipotese['Vento']
    angulo = parametros_hipotese['Ângulo']
    
    # Inicialize o dicionário para armazenar as pressões para cada cabo
    pressao_ven_cabo = {}
    forcas_por_tronco = {}
    pressao_tronco_int = {}
    
    # Loop sobre as fases em geometria_isolador
    for fase, parametros_fase in geometria_isolador.items():
        # Calcule as forças longitudinais e transversais para cada hipótese e fase
        pressao_longitudinal = round(ml.pressao(massa, vento_1)*parametros_fase["gt"]*np.cos(np.radians(angulo)), 2)*1.2
        pressao_trans = round(ml.pressao(massa, vento_1)*parametros_fase["gt"]*np.sin(np.radians(angulo)), 2)*1.2
        
        if 6 <= hipotese <= 10:
            pressao_longitudinal = round(ml.pressao(massa, vento_1)*1*np.cos(np.radians(angulo)), 2)*1.2
            pressao_trans = round(ml.pressao(massa, vento_1)*1*np.sin(np.radians(angulo)), 2)*1.2
            
        
        forcas_longitudinais = round(pressao_longitudinal*isolador["area"], 2)
        forcas_transversais = round(pressao_trans*isolador["area"], 2)
        
        # Armazene as forças no dicionário forcas_por_fase
        forcas_por_fase[fase] = {
            "Forca L": forcas_longitudinais,
            "Forca T": forcas_transversais
        }
        
        
    
    # Armazene as forças para cada hipótese no dicionário forcas_por_hipotese_e_fase
    forca_isolador[hipotese] = forcas_por_fase
    
    for chave, valor in geometria_estrutura.items():
        
        if dbCabos[dbCabos['Cabo'] == valor["cabo"]]["Diametro (mm)"].iloc[0] < 15:
            Cx_cabo = 1.2
        
        
        pressao_vante = ml.pressao_cabo(gl["vante"], valor["gc_vante"], ml.pressao(massa, vento_1), angulo)*Cx_cabo
        pressao_re = ml.pressao_cabo(gl["re"], valor["gc_re"], ml.pressao(massa, vento_1), angulo)*Cx_cabo
        
        
        if 6 <= hipotese <= 10:
            pressao_vante = ml.pressao_cabo(1, 1, ml.pressao(massa, vento_1), angulo)
            pressao_re = ml.pressao_cabo(1, 1, ml.pressao(massa, vento_1), angulo)  
        
        pressao_ven_cabo[chave] = {
            "re": pressao_re,
            "vante": pressao_vante,
        }
    
    # Armazene as pressões para cada hipótese no dicionário pressao_vento_cabo
    pressao_vento_cabo[hipotese] = pressao_ven_cabo
    
    for i, (chave, valor) in enumerate(tronco.items()):
      if i < len(tronco) - 1:
        pressao_longitudinal_1 = round(ml.pressao(massa, vento_1)*valor["gt"]*np.cos(np.radians(angulo)), 2)*Cxposte
        pressao_trans_1 = round(ml.pressao(massa, vento_1)*valor["gt"]*np.sin(np.radians(angulo)), 2)*Cxposte
        
        if 6 <= hipotese <= 10:
            pressao_longitudinal_1 = round(ml.pressao(massa, vento_1)*1*np.cos(np.radians(angulo)), 2)*Cxposte
            pressao_trans_1 = round(ml.pressao(massa, vento_1)*1*np.sin(np.radians(angulo)), 2)*Cxposte
            
        
        forcas_longitudinais = round(pressao_longitudinal_1*valor["secaoA"]*L_tronco, 2)
        forcas_transversais = round(pressao_trans_1*valor["secaoB"]*L_tronco, 2)
        
        # Armazene as forças no dicionário forcas_por_fase
        forcas_por_tronco[chave] = {
            "Forca L": int(round(forcas_longitudinais,0)),
            "Forca T": int(round(forcas_transversais,0))
        }
        
        pressao_tronco_int[chave] = {
           "L": round(pressao_longitudinal_1,2),
           "T": round(pressao_trans_1,0) 
        }
    
    # Armazene as forças para cada hipótese no dicionário forcas_por_hipotese_e_fase
    forca_poste[hipotese] = forcas_por_tronco
    pressao_tronco[hipotese] = pressao_tronco_int
      
for hipotese, parametros_hipotese in hipoteses.items():
    # Inicialize o dicionário para armazenar as forças para cada fase
    forcas_por_fase = {}
    tracao_cabo = {}
    temp = parametros_hipotese["Temperatura"]

    
    for chave, valor in geometria_estrutura.items():
        cabo_selecionado = dbCabos.loc[dbCabos['Cabo'] == valor["cabo"]]
        
        E = float(cabo_selecionado["E"].iloc[0])
        S = float(cabo_selecionado["S"].iloc[0])
        alpha = float(cabo_selecionado["COEF"].iloc[0])
        p1 = float(cabo_selecionado["Peso (kgf/m)"].iloc[0])
        t1 = temperatura["EDS"]
        t2 = temp
        
        diametro_mm = float(cabo_selecionado["Diametro (mm)"].iloc[0])
        
        if "Cabo_PR_1" in chave or "Cabo_PR_2" in chave:
            T0_vante = tracao_eds_pr["vante"]
            T0_re = tracao_eds_pr["re"]
            
            if "Cabo_PR_2" in chave:
                T0_vante = tracao_eds_pr2["vante"]
                T0_re = tracao_eds_pr2["re"]
            
        else:
            T0_vante = tracao_eds["vante"]
            T0_re = tracao_eds["re"]
        
        
        
        p2_re = np.sqrt(p1**2 + (0.001 * pressao_vento_cabo[hipotese][chave]["re"] * diametro_mm)**2)
        p2_vante = np.sqrt(p1**2 + (0.001 * pressao_vento_cabo[hipotese][chave]["vante"] * diametro_mm)**2)
        
        if 6 <= hipotese <= 10:
            p2_re = np.sqrt(p1**2 + (0.001*0.25 * pressao_vento_cabo[hipotese][chave]["re"] * diametro_mm)**2)
            p2_vante = np.sqrt(p1**2 + (0.001*0.25 * pressao_vento_cabo[hipotese][chave]["vante"] * diametro_mm)**2) 
        
        tracao_cabo[chave] = {
            "re": ml.mudanca_estado(E, S, vao_regulador["re"], alpha, p1, p2_re, t1, t2, T0_re),
            "vante": ml.mudanca_estado(E, S, vao_regulador["vante"], alpha, p1, p2_vante, t1, t2, T0_vante),  
        }
    
    tracao_cabo1[hipotese] = tracao_cabo
    
    

altura_maxima = -float('inf')
altura_maxima_condutor = -float('inf')
cabo_condutor_mais_alto = None
cabo_pr_mais_alto = None

#calculo da árvore de carga

for hipotese, parametros_hipotese in hipoteses.items():
     carga = {}
     carga_homens = 200
     
     if hipotese == 12: # Identificar cabo mais alto
             for chave, valor in geometria_estrutura.items():
                 if chave.startswith("Cabo_PR"):
                     altura_cabo_pr = max(valor['H_re'], valor['H_vante'])
                     if altura_cabo_pr > altura_maxima:
                         altura_maxima = altura_cabo_pr
                         cabo_pr_mais_alto = chave
                 else:
                     altura_cabo_condutor = max(valor['H_re'], valor['H_vante'])
                     if altura_cabo_condutor > altura_maxima_condutor:
                         altura_maxima_condutor = altura_cabo_condutor
                         cabo_condutor_mais_alto = chave
                         
     for chave, valor in geometria_estrutura.items():
         
         cabo_selecionado = dbCabos.loc[dbCabos['Cabo'] == valor["cabo"]]
         peso = float(cabo_selecionado["Peso (kgf/m)"].iloc[0])
         
        
         try:
             T_isolador  = forca_isolador[hipotese][chave]["Forca T"]
             L_isolador = forca_isolador[hipotese][chave]["Forca L"]
             V_isolador = isolador["peso"]
         except KeyError:
            T_isolador = 0
            L_isolador = 0
            V_isolador = 0
          
         Ac_vante = pressao_vento_cabo[hipotese][chave]["vante"]* (diametro_mm /1000) * vao["vante"] 
         Ac_re = pressao_vento_cabo[hipotese][chave]["re"]* (diametro_mm /1000) * vao["re"]
         Ac_T = 0.5*(Ac_vante + Ac_re)*np.cos(np.radians(deflexao/2))
         Ac_L = 0.5*(Ac_vante - Ac_re)*np.cos(np.radians(deflexao/2))
         
         Tc_T =  (tracao_cabo1[hipotese][chave]["re"]+ tracao_cabo1[hipotese][chave]["vante"])*np.sin(np.radians(deflexao/2))
         Tc_L = abs((tracao_cabo1[hipotese][chave]["re"]- tracao_cabo1[hipotese][chave]["vante"])*np.cos(np.radians(deflexao/2)))
  
         T_cabo =    valor["qtd"]*(Ac_T + Tc_T) + T_isolador
         L_cabo =    valor["qtd"]*(Ac_L + Tc_L) + L_isolador
    

         Vmin = valor["qtd"]*vao_de_peso["min"]*peso + V_isolador  
         Vmax = valor["qtd"]*vao_de_peso["max"]*peso + V_isolador
         
         if 1 <= hipotese <= 5:
           
             FS_T = fator_sobrecarga["Vento Máximo"]["T"]
             FS_L = fator_sobrecarga["Vento Máximo"]["L"]
             FS_V_normal = fator_sobrecarga["Vento Máximo"]["V_normal"]
             FS_V_reduzido = fator_sobrecarga["Vento Máximo"]["V_reduzido"]
             
         if 6 <= hipotese <= 10:
             
             Ac_vante = 0.25*pressao_vento_cabo[hipotese][chave]["vante"]* (diametro_mm /1000) * vao["vante"] 
             Ac_re = 0.25*pressao_vento_cabo[hipotese][chave]["re"]* (diametro_mm /1000) * vao["re"]
             Ac_T = 0.5*(Ac_vante + Ac_re)*np.cos(np.radians(deflexao/2))
             Ac_L = 0.5*(Ac_vante - Ac_re)*np.cos(np.radians(deflexao/2))
             
             Tc_T =  (tracao_cabo1[hipotese][chave]["re"]+ tracao_cabo1[hipotese][chave]["vante"])*np.sin(np.radians(deflexao/2))
             Tc_L = abs((tracao_cabo1[hipotese][chave]["re"]- tracao_cabo1[hipotese][chave]["vante"])*np.cos(np.radians(deflexao/2)))
      
             T_cabo =    valor["qtd"]*(Ac_T + Tc_T) + T_isolador
             L_cabo =    valor["qtd"]*(Ac_L + Tc_L) + L_isolador
        

             Vmin = valor["qtd"]*vao_de_peso["min"]*peso + V_isolador  
             Vmax = valor["qtd"]*vao_de_peso["max"]*peso + V_isolador
             
             
         
         
         if hipotese == 11: #construção
         
             Vmax = 1.25*valor["qtd"]*max(tracao_cabo1[hipotese][chave]["re"],tracao_cabo1[hipotese][chave]["vante"])*0.93 
             Vmin = 0
             L_cabo = 0.1132*valor["qtd"]*max(tracao_cabo1[hipotese][chave]["re"],tracao_cabo1[hipotese][chave]["vante"])
             T_cabo = (tracao_cabo1[hipotese][chave]["re"]+ tracao_cabo1[hipotese][chave]["vante"])*np.sin(np.radians(deflexao/2))*valor["qtd"]
             
             FS_T = fator_sobrecarga["Construção/Manutenção"]["T"]
             FS_L = fator_sobrecarga["Construção/Manutenção"]["L"]
             FS_V_normal = fator_sobrecarga["Construção/Manutenção"]["V_normal"]
             FS_V_reduzido = fator_sobrecarga["Construção/Manutenção"]["V_reduzido"]
         
            
         if hipotese == 12: #cabo rompido
         
                     
         
            T_cabo = ((tracao_cabo1[hipotese][chave]["re"]+ tracao_cabo1[hipotese][chave]["vante"])*np.sin(np.radians(deflexao/2))+ Ac_T)*valor["qtd"]
            L_cabo = (tracao_cabo1[hipotese][chave]["re"]- tracao_cabo1[hipotese][chave]["vante"])*np.cos(np.radians(deflexao/2))*valor["qtd"]
  
            Vmin = valor["qtd"]*vao_de_peso["min"]*peso       
            Vmax = valor["qtd"]*vao_de_peso["max"]*peso    
                    
         
            if chave == cabo_condutor_mais_alto:
                 Vmax = Vmax*0.7
                 Vmin = Vmin*0.7
                 T_cabo = (tracao_cabo1[hipotese][chave]["re"])*np.sin(np.radians(deflexao/2))*valor["qtd"]
                 L_cabo = 1*(tracao_cabo1[hipotese][chave]["re"])*np.cos(np.radians(deflexao/2))*valor["qtd"]
                 
            FS_T = fator_sobrecarga["Ruptura de cabo"]["T"]
            FS_L = fator_sobrecarga["Ruptura de cabo"]["L"]
            FS_V_normal = fator_sobrecarga["Ruptura de cabo"]["V_normal"]
            FS_V_reduzido = fator_sobrecarga["Ruptura de cabo"]["V_reduzido"]
                 
         if hipotese ==13: #cabo rompido
             T_cabo = ((tracao_cabo1[hipotese][chave]["re"]+ tracao_cabo1[hipotese][chave]["vante"])*np.sin(np.radians(deflexao/2))+ Ac_T)*valor["qtd"]
             L_cabo = (tracao_cabo1[hipotese][chave]["re"]- tracao_cabo1[hipotese][chave]["vante"])*np.cos(np.radians(deflexao/2))*valor["qtd"]
  
             Vmin = valor["qtd"]*vao_de_peso["min"]*peso       
             Vmax = valor["qtd"]*vao_de_peso["max"]*peso    
                    
                    
             if chave == cabo_condutor_mais_alto:
                 Vmax = Vmax*0.7
                 Vmin = Vmin*0.7
                 T_cabo = (tracao_cabo1[hipotese][chave]["vante"])*np.sin(np.radians(deflexao/2))*valor["qtd"]
                 L_cabo = 1*(tracao_cabo1[hipotese][chave]["vante"])*np.cos(np.radians(deflexao/2))*valor["qtd"]
                 
             FS_T = fator_sobrecarga["Ruptura de cabo"]["T"]
             FS_L = fator_sobrecarga["Ruptura de cabo"]["L"]
             FS_V_normal = fator_sobrecarga["Ruptura de cabo"]["V_normal"]
             FS_V_reduzido = fator_sobrecarga["Ruptura de cabo"]["V_reduzido"]

         if hipotese ==14: # para-raio rompido
         

                         
             if chave == cabo_pr_mais_alto:
                 Vmax = Vmax*0.7
                 Vmin = Vmin*0.7
                 T_cabo = (tracao_cabo1[hipotese][chave]["re"])*np.sin(np.radians(deflexao/2)+ Ac_T)*valor["qtd"]+V_isolador
                 L_cabo = 1.25*(tracao_cabo1[hipotese][chave]["re"])*np.cos(np.radians(deflexao/2))*valor["qtd"]+L_isolador
                 
             FS_T = fator_sobrecarga["Ruptura de cabo"]["T"]
             FS_L = fator_sobrecarga["Ruptura de cabo"]["L"]
             FS_V_normal = fator_sobrecarga["Ruptura de cabo"]["V_normal"]
             FS_V_reduzido = fator_sobrecarga["Ruptura de cabo"]["V_reduzido"]
                 
         if hipotese ==15:  # para-raio rompido
             if chave == cabo_pr_mais_alto:
                 Vmax = Vmax*0.7
                 Vmin = Vmin*0.7
                 T_cabo = (tracao_cabo1[hipotese][chave]["vante"])*np.sin(np.radians(deflexao)+ Ac_T)*valor["qtd"]+V_isolador
                 L_cabo = 1.25*(tracao_cabo1[hipotese][chave]["vante"])*np.cos(np.radians(deflexao))*valor["qtd"]+L_isolador    
                 
             FS_T = fator_sobrecarga["Ruptura de cabo"]["T"]
             FS_L = fator_sobrecarga["Ruptura de cabo"]["L"]
             FS_V_normal = fator_sobrecarga["Ruptura de cabo"]["V_normal"]
             FS_V_reduzido = fator_sobrecarga["Ruptura de cabo"]["V_reduzido"]
                 
         if hipotese == 16:            # Contenção cascata
               
         
            if chave.startswith("Cabo_PR"):
                     L_cabo = max(tracao_cabo1[hipotese][chave]["vante"],tracao_cabo1[hipotese][chave]["re"])*0.60 * valor["qtd"]
            else:
                     L_cabo = max(tracao_cabo1[hipotese][chave]["vante"],tracao_cabo1[hipotese][chave]["re"])*0.40 * valor["qtd"]
                     
            FS_T = fator_sobrecarga["Contenção em cascata"]["T"]
            FS_L = fator_sobrecarga["Contenção em cascata"]["L"]
            FS_V_normal = fator_sobrecarga["Contenção em cascata"]["V_normal"]
            FS_V_reduzido = fator_sobrecarga["Contenção em cascata"]["V_reduzido"]
                     
         if hipotese == 17:  # Temperatura mínima
             T_cabo = ((tracao_cabo1[hipotese][chave]["re"]+ tracao_cabo1[hipotese][chave]["vante"])*np.sin(np.radians(deflexao/2))+
                       Ac_T)*valor["qtd"]+T_isolador
             L_cabo = (tracao_cabo1[hipotese][chave]["re"]- tracao_cabo1[hipotese][chave]["vante"])*np.cos(np.radians(deflexao/2))*valor["qtd"]+L_isolador
             
             Vmin = valor["qtd"]*vao_de_peso["min"]*peso + V_isolador
             
             Vmax = valor["qtd"]*vao_de_peso["max"]*peso + V_isolador
             
             FS_T = fator_sobrecarga["Vento Máximo"]["T"]
             FS_L = fator_sobrecarga["Vento Máximo"]["L"]
             FS_V_normal = fator_sobrecarga["Vento Máximo"]["V_normal"]
             FS_V_reduzido = fator_sobrecarga["Vento Máximo"]["V_reduzido"]
             
         
         if hipotese == 18:  # Terminal
                Ac_vante = pressao_vento_cabo[hipotese][chave]["vante"]* (diametro_mm /1000) * vao["vante"] 
                Ac_re = pressao_vento_cabo[hipotese][chave]["re"]* (diametro_mm /1000) * vao["re"]
                Ac_T = 0.5*(Ac_vante + Ac_re)*np.cos(np.radians(deflexao/2))
                Ac_L = 0.5*(Ac_vante - Ac_re)*np.cos(np.radians(deflexao/2))
                
                Tc_T =  (max(tracao_cabo1[hipotese][chave]["re"],tracao_cabo1[hipotese][chave]["vante"]))*np.sin(np.radians(deflexao/2))
                Tc_L = (max((abs(tracao_cabo1[hipotese][chave]["re"]), abs(tracao_cabo1[hipotese][chave]["vante"])))*np.cos(np.radians(deflexao/2)))
         
                T_cabo =    valor["qtd"]*max(Ac_T, Tc_T) + T_isolador
                L_cabo =    valor["qtd"]*max(Ac_L , Tc_L) + L_isolador
           
    
                Vmin = valor["qtd"]*vao_de_peso["min"]*peso + V_isolador  
                Vmax = valor["qtd"]*vao_de_peso["max"]*peso + V_isolador
            
            
    
         carga[chave] = {
            "T": int(round(T_cabo*FS_T,0)),
            "Vmin":int(round(Vmin*FS_V_reduzido ,0)),
            "Vmax":int(round(Vmax*FS_V_normal ,0)),
            "L": int(round(abs(L_cabo*FS_L),0)),
           }
         
         
     arvore_carga[hipotese] = carga
         
 # CALCULO DOS MOMENTOS

momentoT = {}
momentoL = {}
resultante = {}
ang_resultante = {}
fator_reducao = {}
carga_resultante={}

for chave, hipotese in hipoteses.items():
        momento_t_temp = {}
        mt = 0
        ml = 0
        res = 0
        for fase, valor in geometria_estrutura.items(): 
            mt += valor["H"]*arvore_carga[chave][fase]["T"] + valor["X"]*arvore_carga[chave][fase]["Vmax"]
            ml += valor["H"]*arvore_carga[chave][fase]["L"]
            
        for i, (troncos, dados2) in enumerate(tronco.items()):
             if i < len(tronco) - 1: 
               mt += dados2["centroide"]*forca_poste[chave][troncos]["Forca T"]
               ml += dados2["centroide"]*forca_poste[chave][troncos]["Forca L"]
           
        momentoT[chave] = int(round(mt,0))
        momentoL[chave] = int(round(ml,0))
        
        resultante[chave] = int(round(np.sqrt(momentoT[chave]**2 + momentoL[chave]**2)/(altura_estrutura-0.8),0))
        if poste != "R":
            ang_resultante[chave] = abs(round( deflexao_bissetriz- np.rad2deg(np.arctan( momentoL[chave]/ momentoT[chave])),2))
        if poste == "DT":
            fator_reducao[chave] = round(min(1, 1.0329*np.exp(-0.00722*ang_resultante[chave])),2)
            if ang_resultante[chave] > 85 :
               fator_reducao[chave] = 0.5
        else:
            fator_reducao[chave] = 1
        carga_resultante[chave]= int(round(resultante[chave]/ fator_reducao[chave],0))
            
        




caminho_arquivo = 'arvore_carga.txt'
salvar_em_txt(arvore_carga, caminho_arquivo)
# Criar um DataFrame a partir do dicionário
# Criar o elemento raiz do XML

root = ET.Element("arvore_carga")



# Criar a árvore XML
arvore_xml = ET.ElementTree(root)

# Salvar o arquivo XML
arvore_xml.write("arvore_carga.xml", encoding="utf-8", xml_declaration=True)


def plot_arrows_exemplo(ax, x, y, label_0, label_90, label_45, cor):
  
    ax.arrow(x - 1.15, y, 1, 0, head_width=0.1, head_length=0.1, fc=cor, ec=cor, linewidth=1.2)  # seta a 0 graus  
    ax.arrow(x, y, 0, -1, head_width=0.1, head_length=0.1, fc=cor, ec=cor, linewidth=1.2)  # seta a -90 graus
    ax.arrow(x, y, 1 / np.sqrt(2), 1 / np.sqrt(3), head_width=0.1, head_length=0.1, fc=cor, ec=cor, linewidth=1.2)  # seta a 45 graus
    ax.annotate(f'{label_0}\n', (x - 0.7 , y - 0.6), textcoords="offset points", xytext=(0, 10), ha='center', color = cor)
    ax.annotate(f'{label_90}\n', (x, y - 2), textcoords="offset points", xytext=(0, 0), ha='center', color = cor)
    ax.annotate(f'{label_45}\n', (x + 0.5 + 2 / np.sqrt(2) / 2, y + 0.7 / np.sqrt(2) / 10),textcoords="offset points", xytext=(0, 10), ha='center', color = cor)


# Função para plotar setas em uma determinada posição (x, y) com diferentes direções e valores associados
def plot_arrows(ax, x, y, label_0, label_90, label_45):
    if label_0 != 0 :
        ax.arrow(x - 1.15, y, 1, 0, head_width=0.1, head_length=0.1, fc='black', ec='black', linewidth=1.2)  # seta a 0 graus
     
        
     
    ax.arrow(x, y, 0, -1, head_width=0.1, head_length=0.1, fc='black', ec='black', linewidth=1.2)  # seta a -90 graus
    if label_45 != 0 :
        ax.arrow(x, y, 1 / np.sqrt(2), 1 / np.sqrt(3), head_width=0.1, head_length=0.1, fc='black', ec='black', linewidth=1.2)  # seta a 45 graus

    # Adicionar valores associados às setas
    if label_0 != 0 :
        ax.annotate(f'{label_0}\n', (x - 0.7 , y - 0.6), textcoords="offset points", xytext=(0, 10), ha='center')
     
    linhas = label_90.split('\n')   
    for i, linha in enumerate(linhas):
        ax.annotate(f'{linha}\n', (x+0.03+i*0.3, y - 0.9), textcoords="offset points", xytext=(0, i * 0.0), ha='center', rotation=90)
        
    if label_45 != 0 :
        ax.annotate(f'{label_45}\n', (x + 0.5 + 2 / np.sqrt(2) / 2, y + 0.7 / np.sqrt(2) / 10),
                textcoords="offset points", xytext=(0, 10), ha='center')


# Função para plotar setas em uma determinada posição (x, y) com diferentes direções e valores associados
def plot_arrows2(ax, h_p, valor_T, valor_L):


    if valor_T != 0 :
        ax.arrow(-4, h_p, 1, 0, head_width=0.1, head_length=0.1, fc='blue', ec='blue', linewidth=1.2) 
        ax.annotate(f'{valor_T}\n', (-3.5, h_p-0.7), textcoords="offset points", xytext=(0, 10), ha='center', color='blue')
        
    if valor_L != 0 :
         ax.arrow(-2.9, h_p, 1 / np.sqrt(2), 1 / np.sqrt(3), head_width=0.1, head_length=0.1, fc='blue', ec='blue', linewidth=1.2)
         ax.annotate(f'{valor_L}', (-1.8, h_p + 1 / np.sqrt(2) / 10), textcoords="offset points", xytext=(0, 10), ha='center', color='blue')
                

amostras = list(carga_resultante.keys())
valores = list(carga_resultante.values())

plt.grid(True, linestyle='--', alpha=0.7)
# Criar um gráfico de barras
plt.bar(amostras, valores, color='blue')

# Adicionar uma linha horizontal para o valor da carga inicial
plt.axhline(y=Carga_inicial, color='red', linestyle='-', label='Carga nominal do Poste')

plt.ylim(0, Carga_inicial + 500)
# Adicionar rótulos ao gráfico
plt.title('Gráfico das Cargas finais no poste')
plt.xlabel('Estado de Carga')
plt.ylabel('Cargas finais (kgf)')

# Adicionar uma legenda
plt.legend()

# Definir a escala no eixo x para mostrar todas as amostras
plt.xticks(amostras)
plt.savefig('arvore/resumo.png', dpi=300, bbox_inches='tight',  format='png')    



for chave, hipotese in hipoteses.items():
    # Criar uma figura com DPI de 300
    fig, ax = plt.subplots(dpi=300,figsize=(10, 10))
    vertices = [
    (-0.0005 * tronco[5]["Face_A"], 0),
    (-0.0005 * tronco[1]["Face_A"], altura_estrutura),
    (0.0005 * tronco[1]["Face_A"], altura_estrutura),
    (0.0005 * tronco[5]["Face_A"], 0)
]
    quadrado = Polygon(vertices, closed=True, edgecolor='black', facecolor='#D3D3D3')
    ax.add_patch(quadrado)
    
    # Definir limites do gráfico
    altura = altura_estrutura
    ax.set_xlim(-9, 8)
    ax.set_ylim(1, altura+2)
    
    # Adicionar grades para melhor orientação
    ax.grid(True, which='both', linestyle='--', linewidth=0.5)
    ax.minorticks_on()
    ax.grid(True, which='minor', linestyle=':', linewidth=0.5)


    # Plotar setas em diferentes posições com valores associados
    for fase, dados in geometria_estrutura.items():
        x = dados['X']
        y = dados['H']
        carga_info = arvore_carga[chave][fase]  # Usei a carga do nível 1 para simplificar, ajuste conforme necessário
        plot_arrows(ax, x, y, carga_info["T"], f" {carga_info['Vmax']}\n({carga_info['Vmin']})", carga_info["L"])
    
    
    for i, (troncos, dados2) in enumerate(tronco.items()):
     if i < len(tronco) - 1: 
      
        hc = dados2["centroide"]
        carga_info2 = forca_poste[chave][troncos]
        plot_arrows2(ax,hc, carga_info2["Forca T"],carga_info2["Forca L"])
    
    
    x = 6
    y = 7.5
    plot_arrows_exemplo(ax, x, y, "T","V", "L", "black")
    ax.text(4.5,9, "Cargas nos cabos", va='center', fontsize=10, color='black')
    
    x = 6
    y = 3.8
    plot_arrows_exemplo(ax, x, y, "T","V", "L", "blue")
    ax.text(4.5,5.2, "Cargas no poste", va='center', fontsize=10, color='blue')
    
    chaves = hipoteses.keys()
    
    # Convertendo as chaves para uma lista (opcional, dependendo do que você precisa)
    lista_de_chaves = chave
    cor_Texto = "purple"
    inicio  = - 8
    plt.title("Hipótese de carga: " + str(hipotese["nome"]))
    plt.xlabel('Eixo Transversal')
    plt.ylabel('Altura (m)')
    ax.text(inicio, altura +1.6, 'Estrutura:' +nome_estrutura , ha='left', va='center', fontsize=10, color='purple')
    ax.text(inicio, altura +1.1, "Deflexão = " + str(deflexao) + "°", ha='left', va='center', fontsize=10, color='purple')
    ax.text(inicio, altura +0.6, 'Cargas em kgf', ha='left', va='center', fontsize=10, color='purple')
    ax.text(inicio, altura +0.1, 'Vento = ' + str(hipotese["Vento"]) +"m/s; Ângulo = " +str(hipotese["Ângulo"])+"°", ha='left', va='center', fontsize=10, color='purple')
    ax.text(inicio, altura - 0.4, "Mom. Long. = " + str(momentoL[chave]) +"kgfm", ha='left', va='center', fontsize=10, color='purple')
    ax.text(inicio, altura -0.9, "Mom. Trans. = " + str(momentoT[chave])+"kgfm", ha='left', va='center', fontsize=10, color='purple')
    ax.text(inicio, altura -1.4, "Carga solicitada = "+ str(resultante[chave]) + "kgf", ha='left', va='center', fontsize=10, color='purple')
    
    if poste =="DT":
        ax.text(inicio, altura -1.9, "Ângulo result. P/ face 'B' = "+ str(ang_resultante[chave]) +"°", ha='left', va='center', fontsize=10, color='purple')
        ax.text(inicio, altura -2.4, "Fator de correção = "+ str(fator_reducao[chave]), ha='left', va='center', fontsize=10, color='purple')
        ax.text(inicio, altura -2.9, "Carga final = "+ str(carga_resultante[chave]) +"kgf", ha='left', va='center', fontsize=10, color='purple')
    if poste =="R":   
        ax.text(inicio, altura -1.9, "Fator de correção = "+ str(fator_reducao[chave]), ha='left', va='center', fontsize=10, color='purple')
        ax.text(inicio, altura -2.4, "Carga final = "+ str(carga_resultante[chave]) +"kgf", ha='left', va='center', fontsize=10, color='purple')
    

    
    fig.savefig("arvore/hipotese de carga" + str(lista_de_chaves)+".png", bbox_inches='tight',  format='png')
    fig.savefig("arvore/hipotese de carga" + str(lista_de_chaves)+".svg", bbox_inches='tight',  format='svg')
    # Exibir o gráfico
    #plt.show()




# Exibir o gráfico
plt.show()

def adicionar_tabela(doc, dados):
    # Obter o número de linhas e colunas a partir dos dados
    num_linhas = len(dados)
    num_colunas = len(dados[0])

    # Adicione a tabela diretamente ao documento com base no número de linhas e colunas dos dados
    table = doc.add_table(rows=num_linhas, cols=num_colunas)
    table.style = "Grid Table 4 Accent 1"

    # Preencha os dados da tabela
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            cell.text = str(dados[row_index][cell_index])
            
def inserir_grafico(doc, imagem_path, titulo, estilo_titulo=None, largura=Inches(6)):
    # Insere um parágrafo com o título do gráfico
    titulo_paragrafo = doc.add_paragraph()
    titulo_run = titulo_paragrafo.add_run(titulo)
    
    # Aplicar estilo de caractere ao título, se fornecido
    if estilo_titulo:
        titulo_run.font.name = estilo_titulo['font_name']
        titulo_run.font.size = estilo_titulo['font_size']
        titulo_run.font.color.rgb = estilo_titulo['font_color']
    
    titulo_paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Insere a imagem do gráfico
    imagem_paragrafo = doc.add_paragraph()
    run = imagem_paragrafo.add_run()
    run.add_picture(imagem_path, width=largura)
    imagem_paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adicione uma quebra de linha após a imagem
    doc.add_paragraph()

#fazer o memorial
# Crie um documento do Word
doc = Document("modelo.docx")
styles = doc.styles
    
# Defina as palavras que você deseja procurar e substituir

referencias = {
    "Numero_desenho": "12343",
    "Nome_documento": "Capacidade operativa",
    "nome_lt": nome_LT,
    }

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
doc.add_heading("OBJETIVO", level = 1)

data = "Este relatório tem como objetivo apresentar o cálculo estrutural para as estruturas da " + nome_LT + "."
doc.add_paragraph(data, style="First Paragraph")

doc.add_heading("DADOS BÁSICOS", level = 1)
doc.add_heading("Cabos", level = 2)

dados = [
    ["Característica","Unidade", "Condutor", "Cabo guarda",],
    ["Nome","", dadosCabo["Cabo"].iloc[0],dados_pr1["Cabo"].iloc[0]],
    ["Seção","mm²", dadosCabo["Seção (mm²)"].iloc[0],dados_pr1["Seção (mm²)"].iloc[0]],
    ["Diâmetro","mm", dadosCabo["Diametro (mm)"].iloc[0],dados_pr1["Diametro (mm)"].iloc[0]],
    ["Peso","kgf/m", dadosCabo["Peso (kgf/m)"].iloc[0],dados_pr1["Peso (kgf/m)"].iloc[0]],
    ["Mód. Elast. final","kgf/mm²", dadosCabo["E"].iloc[0],dados_pr1["E"].iloc[0]],
    ["Coef Dilat. Térm. final","1/°C", dadosCabo["COEF"].iloc[0],dados_pr1["COEF"].iloc[0]],
    ]

adicionar_tabela(doc,dados)

doc.add_heading("Cadeia de Isoladores", level = 2)
dados = [
    ["Área Exposta (mm²)", "Peso (kgf)"],
    [isolador["area"], isolador["peso"]]
    ]

adicionar_tabela(doc,dados)

massa = round(massa,3)
doc.add_heading("Ação do Vento", level = 2)
data = "As velocidades de vento a serem utilizadas no projeto da linha e as respectivas pressões e cargas atuantes nos cabos, isoladores e estruturas serão calculadas de acordo com a metodologia da publicação IEC 60826 - International Eletrotechnical Comission: Design Criteria of Overhead Transmission Lines, considerando-se: "
doc.add_paragraph(data, style="First Paragraph")
doc.add_paragraph(" - Terreno com rugosidade categoria " + rugosidade + ";", style="First Paragraph")
doc.add_paragraph(" - Velocidade de Vento de Referência " + str(vento["projeto"]) + " m/s;" , style="First Paragraph")
doc.add_paragraph(" - Velocidade de Vento de Alta intensidade " + str(vento["extremo"]) + " m/s;" , style="First Paragraph")
doc.add_paragraph(f" - Período de retorno T = {periodo_retorno} anos;", style="First Paragraph")
doc.add_paragraph(f" - ρ = massa específica do ar = {massa} (temperatura = {temperatura['coincidente']} °C, altitude = {altitude} m);", style="First Paragraph")
doc.add_paragraph(" - g = aceleração da gravidade = 9.81 m/s².",  style="First Paragraph")

doc.add_heading("Ação de vento na estrutura", level = 3)
doc.add_heading("Vento máximo", level = 5)
doc.add_paragraph("A carga de vento máximo atuante na estrutura será determinada de acordo com o prescrito na Publicação IEC 60826, utilizando-se a expressão abaixo:", style="First Paragraph")

if poste == "DT":
    doc.add_paragraph("A_tc = q₀ Cx Gt a L",  style="First Paragraph")
    doc.add_paragraph("Onde:",  style="First Paragraph")
    doc.add_paragraph("- A_tc   é a força exercida pelo vento na estrutura;",  style="First Paragraph")
    doc.add_paragraph("- q₀   é a pressão dinâmica;",  style="First Paragraph")
    doc.add_paragraph("- Cx   é o coeficiente de arrasto;",  style="First Paragraph")
    doc.add_paragraph("- Gt   é o fator combinado de vento para o suporte;",  style="First Paragraph")
    doc.add_paragraph("- a   é a dimensão do tronco na face de incidência do vento ao nível do centro geométrico do elemento;",  style="First Paragraph")
    doc.add_paragraph("- L   é o comprimento do tronco.",  style="First Paragraph")

    ## colocar aqui se for R
    
doc.add_heading("Vento de alta intensidade", level = 5)    
doc.add_paragraph("As cargas de vento de alta intensidade são calculadas de forma similar ao item anterior, utilizando-se a expressão abaixo:",  style="First Paragraph")
doc.add_paragraph("A_tc = qi Cx a L",  style="First Paragraph")
doc.add_paragraph("Onde:",  style="First Paragraph")
doc.add_paragraph("qi   É a pressão dinâmica para vento de alta intensidade.",  style="First Paragraph")

doc.add_heading("Ação do vento nas cadeias de isoladores", level = 3)  
doc.add_heading("Vento máximo", level = 5)  
doc.add_paragraph("Fi = q₀ Gi Cxi Ai, onde: ",  style="First Paragraph")  
doc.add_paragraph("- Fi   Carga de vento, atuante na direção do vento, em kgf;",  style="First Paragraph")
doc.add_paragraph("- q₀   Pressão dinâmica;",  style="First Paragraph")
doc.add_paragraph("- Gi   Fator de rajada, obtido da figura 5 da Publicação IEC 60826;",  style="First Paragraph")
doc.add_paragraph("- Cx   é o coeficiente de arrasto = 1.2;",  style="First Paragraph")
doc.add_paragraph("- Ai   Área exposta ao vento;",  style="First Paragraph")

doc.add_heading("Vento de alta intensidade", level = 5)  
doc.add_paragraph("Fi = qi Cxi Ai, onde: ",  style="First Paragraph")  
doc.add_paragraph("- Fi   Carga de vento, atuante na direção do vento, em kgf;",  style="First Paragraph")
doc.add_paragraph("- qi   Pressão dinâmica para vento de alta intensidade;",  style="First Paragraph")
doc.add_paragraph("- Cxi   é o coeficiente de arrasto = 1.2;",  style="First Paragraph")
doc.add_paragraph("- Ai   Área exposta ao vento;",  style="First Paragraph")
 

doc.add_heading("Ação do vento nos cabos", level = 3)     
doc.add_heading("Vento máximo", level = 5) 
doc.add_paragraph("Fc = q₀ Gc Gl Cxc Φ L sen²(Ω), onde: ",  style="First Paragraph")  
doc.add_paragraph("- Fc  Carga de vento, atuando na direção perpendicular ao cabo; ",  style="First Paragraph")  
doc.add_paragraph("- Gc   Fator de rajada, obtido da figura 3 da IEC 60826;",  style="First Paragraph")
doc.add_paragraph("- Gl   Fator de vão, conforme figura 4 da IEC 60826;",  style="First Paragraph")
doc.add_paragraph("- Cxc   Coeficiente de arrasto;",  style="First Paragraph")
doc.add_paragraph("- Φ   Diâmetro do cabo;",  style="First Paragraph")
doc.add_paragraph("- L   Vâo médio da estrutura;",  style="First Paragraph")
doc.add_paragraph("- Ω   ângulo entre direção de incidência do vento e o cabo, conforme figura 6 da IEC 60826;",  style="First Paragraph")


doc.add_heading("Vento alta intensidade", level = 5) 
doc.add_paragraph("Fc = qi Cxc Φ L sen²(Ω), onde: ",  style="First Paragraph")  
doc.add_paragraph("- Fc  Carga de vento, atuando na direção perpendicular ao cabo: ",  style="First Paragraph")  
doc.add_paragraph("- qi   Pressão dinâmica para vento de alta intensidade;",  style="First Paragraph")
doc.add_paragraph("- Cxc   Coeficiente de arrasto;",  style="First Paragraph")
doc.add_paragraph("- Φ   Diâmetro do cabo;",  style="First Paragraph")
doc.add_paragraph("- L   Vâo médio da estrutura;",  style="First Paragraph")
doc.add_paragraph("- Ω   ângulo entre direção de incidência do vento e o cabo, conforme figura 6 da IEC 60826;",  style="First Paragraph")

doc.add_heading("Trações nos cabos", level = 2) 
doc.add_heading("Estados de carga", level = 3) 


dados = [["N°", "Descrição", "Temperatura (°C)"],]

for chave, hipotese in hipoteses.items(): 
    dados.append([str(chave), hipotese["nome"], str(hipotese["Temperatura"])])
    
    
adicionar_tabela(doc,dados)
doc.add_heading("Condições básicas para os cálculos", level = 3) 
doc.add_paragraph(f" Vão básico de referência ré: {vao_regulador['re']} m.", style="First Paragraph")
doc.add_paragraph(f" Vão básico de referência Vante: {vao_regulador['vante']} m.", style="First Paragraph")
doc.add_paragraph(f" Vão de peso máximo: {vao_de_peso['max']} m.", style="First Paragraph")
doc.add_paragraph(f" Vão de peso mínimo: {vao_de_peso['min']} m.", style="First Paragraph")
doc.add_paragraph(f" Vão a vante: {vao['vante']} m.", style="First Paragraph")
doc.add_paragraph(f" Vão a ré: {vao['re']} m.", style="First Paragraph")

#doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)





doc.add_heading("Condições de tracionamento dos cabos", level = 3)

dados = [
    ["Cabo", "Tração Final EDS Ré (kgf)", "Tração Final EDS Vante (kgf)"],
    ["Condutor", tracao_eds["re"], tracao_eds["vante"]],
    ["Pára-raios", tracao_eds_pr["re"], tracao_eds_pr["vante"]]
    ]    
adicionar_tabela(doc,dados)   
 

doc.add_heading("Posições dos cabos", level = 3) 
# Preparar os dados para a tabela
dados = [["N°", "Tipo de conexão", "Altura na estrutura (m)", "Altura Cabos Vante (m)", "Altura Cabos Ré (m)", "Gc Vante ", "Gc Ré "]]

for chave, fase in geometria_estrutura.items():
    dados.append([
        str(chave),
        fase.get('Tipo', ''),
        str(fase.get('H', '')),
        str(fase.get('H_vante', '')),
        str(fase.get('H_re', '')),
        str(round(fase.get('gc_vante', ''),2)),
        str(round(fase.get('gc_re', ''),2))
    ])
   
adicionar_tabela(doc,dados)  

doc.add_heading("Pressões nos cabos", level = 3) 

dados = [["Estados de carga", "Fase", "Ré (kgf/m²)", "Vante (kgf/m²)"]]

for numero, fases in pressao_vento_cabo.items():
    for fase, valores in fases.items():
        dados.append([
            str(numero),
            fase,
            str(round(valores['re'], 2)),
            str(round(valores['vante'], 2))
        ])

adicionar_tabela(doc,dados)  


doc.add_heading("Cargas finais nos cabos", level = 2) 

dados = [["Estados de carga", "Fase", "Ré  (kgf)", "Vante (kgf)"]]


for numero, fases in tracao_cabo1.items():
    for fase, valores in fases.items():
        dados.append([
            str(numero),
            fase,
            str(int(valores.get('re', ''))),
            str(int(valores.get('vante', '')))
        ])

adicionar_tabela(doc,dados) 

doc.add_heading("Pressão de vento na estrutura", level = 2) 

dados = [["Estado de carga","Tronco", "Altura do Centroide (m)", "Gt", "Pressão Transversal (kgf/m²)", "Pressão Longitudinal (kgf/m²)"]]

for chave, hipotese in hipoteses.items():
    for i, (numero, valores) in enumerate(tronco.items()):
     if i < len(tronco) - 1: 
            dados.append([
                str(chave),
                str(numero),
                str(round(valores["centroide"],2)),
                str(round(valores["gt"],2)),
                str(round(pressao_tronco[chave][numero]["T"])),
                str(round(pressao_tronco[chave][numero]["L"])),
            ])

adicionar_tabela(doc,dados) 


doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


descricao_hipoteses = { 
    1:"Cargas decorrentes da ação do vento máximo de projeto, com direção transversal, sobre cabos, cadeias de isoladores e estrutura; verticais normais e peso próprio da estrutura. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo).",
    2:"Cargas decorrentes da ação do vento máximo de projeto, com direção longitudinal, sobre cabos, cadeias de isoladores e estrutura; verticais normais e peso próprio da estrutura. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo).",
    3:"Cargas decorrentes da ação do vento máximo de projeto, a 45° com a direção da linha, sobre cabos, cadeias de isoladores e estrutura; verticais normais e peso próprio da estrutura. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo).",
    4:"Cargas decorrentes da ação do vento máximo de projeto, a 60° com a direção da linha, sobre cabos, cadeias de isoladores e estrutura; verticais normais e peso próprio da estrutura. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo).",
    5:"Cargas decorrentes da ação do vento máximo de projeto, a 75° com a direção da linha, sobre cabos, cadeias de isoladores e estrutura; verticais normais e peso próprio da estrutura. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo).",
    6:"Cargas decorrentes da ação do vento de alta intensidade, com direção transversal, sobre cabos, cadeias de isoladores e estrutura; verticais normais e peso próprio da estrutura. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo).",
    7:"Cargas decorrentes da ação do vento de alta intensidade, com direção longitudinal, sobre cabos, cadeias de isoladores e estrutura; verticais normais e peso próprio da estrutura. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo).",
    8:"Cargas decorrentes da ação do vento de alta intensidade, a 45° com a direção da linha, sobre cabos, cadeias de isoladores e estrutura; verticais normais e peso próprio da estrutura. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo).", 
    9:"Cargas decorrentes da ação do vento de alta intensidade, a 60° com a direção da linha, sobre cabos, cadeias de isoladores e estrutura; verticais normais e peso próprio da estrutura. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo).", 
    10:"Cargas decorrentes da ação do vento de alta intensidade, a 75° com a direção da linha, sobre cabos, cadeias de isoladores e estrutura; verticais normais e peso próprio da estrutura. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo).", 
    11: "Cargas de construção/manutenção atuando simultaneamente em qualquer combinação possível de para-raios  ou em qualquer combinação possível de fases, com apenas um dos cabos em lançamento e os demais já lançados; peso próprio da estrutura; sem vento. A carga vertical máxima deverá ser determinada considerando-se carga vertical proveniente de cabo ancorado ao solo, com ângulo de 70° e tração de 125% da tração EDS.", 
    12:"Carga longitudinal correspondente a 100% da tração EDS atuando em qualquer uma das fases no sentido ré; verticais normais e peso próprio da estrutura; sem vento. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo). ",
    13: "Carga longitudinal correspondente a 100% da tração EDS atuando em qualquer uma das fases no sentido Vante; verticais normais e peso próprio da estrutura; sem vento. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo). ",
    14: "Carga longitudinal correspondente a 125% da tração EDS atuando em qualquer cabo para-raios  no sentido ré; verticais normais e peso próprio da estrutura; sem vento. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo). ", 
    15: "Carga longitudinal correspondente a 125% da tração EDS atuando em qualquer cabo para-raios  no sentido Vante; verticais normais e peso próprio da estrutura; sem vento. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo). ",
    16: "Cargas longitudinais nos para-raios  correspondentes a 60% das trações EDS e cargas longitudinais nas fases correspondentes a 40% das trações EDS atuando simultaneamente em todos os cabos; verticais normais e peso próprio da estrutura; sem vento. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo).",    
    17: "Cargas dos cabos na temperatura mínima, sem vento, a estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo). ",
    18: "Cargas decorrentes da ação do vento máximo de projeto, com direção transversal, sobre cabos, cadeias de isoladores e estrutura; verticais normais e peso próprio da estrutura considerando cabos para o lado de maior tracionamento. A estrutura deve ser também verificada para verticais reduzidas (estrutura com vão gravante mínimo).",
    }

doc.add_heading("HIPÓTESES PARA DIMENSIONAMENTO ESTRUTURAL", level = 1) 


for numero, hipotese in hipoteses.items():
    doc.add_heading(f"{hipotese['nome']}", level = 2) 
    doc.add_paragraph(f"{descricao_hipoteses[numero]}" ,style="First Paragraph")

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

doc.add_heading("FATORES DE SOBRECARGA", level = 1) 

doc.add_paragraph("As cargas atuantes em cada uma das hipóteses de carga indicadas nos itens anteriores devem ser majoradas pelos seguintes fatores de sobrecarga:" ,style="First Paragraph")
dados = [["Estados de carga", "Transversal", "Longitudinal", "Vertical Máximo", "Vertical Mínimo"]]

for numero, valores in fator_sobrecarga.items():
        dados.append([
            numero,
            str(valores.get('T', '')),
            str(valores.get('L', '')),
            str(valores.get('V_normal', '')),
            str(valores.get('V_reduzido', '')),
        ])

adicionar_tabela(doc,dados) 

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

doc.add_heading("CARREGAMENTOS", level = 1) 

for chave, hipotese in hipoteses.items(): 
    doc.add_heading(f"{hipotese['nome']}", level = 2) 
    inserir_grafico(doc, f'arvore/hipotese de carga{chave}.png', "", largura=Inches(6.89))
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    

doc.add_heading("RESUMO", level = 1) 
inserir_grafico(doc, 'arvore/resumo.png', "", largura=Inches(6.89))
    
 # Ajustar as tabelas
for tabela_index, tabela in enumerate(doc.tables[1:], start=1):
     # Percorra as linhas da tabela
     for row in tabela.rows:
         # Defina a altura da linha
         row.height = Cm(0.6)  # Altere o valor para a altura desejada
         # Alinhe o conteúdo verticalmente
         
         for cell in row.cells:
             cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER   
    

doc.save('arvore/'+str(numero_documento)+ nome_estrutura.replace("/", "-") + '.docx')
# Registre o tempo de término
tempo_final = time.time()

# Calcule o tempo decorrido
tempo_decorrido = tempo_final - tempo_inicial

print(f"Tempo decorrido: {tempo_decorrido} segundos")